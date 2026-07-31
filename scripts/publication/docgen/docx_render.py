"""DOCX renderer for the DarkSage publication pipeline.

Applies the DSF-003 visual design system directly via python-docx +
low-level OOXML manipulation (python-docx has no high-level API for page
background fill, field codes, or image alt text, so those are built by hand
against the OOXML schema here). Deterministic: no wall-clock timestamp is
ever written; the only "generated" timestamp is the caller-supplied git
baseline commit (see generate.py).
"""

from __future__ import annotations

import zipfile
from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from . import design
from .diagrams import DiagramRegistry
from .inline import Span, parse_inline
from .model import (
    Blockquote,
    CodeBlock,
    Document as ModelDocument,
    Figure,
    Heading,
    ListBlock,
    Paragraph,
    Rule,
    Table,
)

_STATUS_RGB = {
    "signal_green": design.SIGNAL_GREEN,
    "warning_amber": design.WARNING_AMBER,
    "soft_gray": design.SOFT_GRAY,
}


def _rgb(hex6: str) -> RGBColor:
    return RGBColor.from_string(hex6)


# ---------------------------------------------------------------------------
# Low-level OOXML helpers (python-docx has no public API for these)
# ---------------------------------------------------------------------------


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _table_borders(table, hex_color: str, sz: str = "4") -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_color)
        borders.append(el)
    tblPr.append(borders)


def _table_no_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _fixed_layout(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def _add_field(paragraph, instr: str, cached_result: str = "") -> None:
    """Insert a Word complex field (PAGE, TOC, ...). Word/LibreOffice compute
    the real value when the field is updated; cached_result is shown until
    then (deterministic placeholder, never wall-clock derived)."""
    r1 = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    r1._r.append(fld1)

    r2 = paragraph.add_run()
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    r2._r.append(instr_el)

    r3 = paragraph.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "separate")
    r3._r.append(fld2)

    r4 = paragraph.add_run(cached_result)

    r5 = paragraph.add_run()
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "end")
    r5._r.append(fld3)


def _set_alt_text(inline_shape, description: str) -> None:
    docPr = inline_shape._inline.find(qn("wp:docPr"))
    if docPr is not None:
        docPr.set("descr", description)
        docPr.set("title", description[:60])


def _paragraph_bottom_rule(paragraph, hex_color: str, sz: str = "18") -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _unlink_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for p in list(section.header.paragraphs):
        p.text = ""
    for p in list(section.footer.paragraphs):
        p.text = ""


# ---------------------------------------------------------------------------
# Run/paragraph writers
# ---------------------------------------------------------------------------


def _write_spans(paragraph, spans: List[Span], base_size: float, color_hex: Optional[str] = None) -> None:
    for span in spans:
        run = paragraph.add_run(span.text)
        run.bold = span.bold
        run.italic = span.italic
        if span.code:
            run.font.name = design.MONO_FONT
            run.font.size = Pt(base_size - 0.5)
        else:
            run.font.name = design.BODY_FONT
            run.font.size = Pt(base_size)
        if span.status_color:
            run.font.color.rgb = _rgb(_STATUS_RGB[span.status_color])
        elif color_hex:
            run.font.color.rgb = _rgb(color_hex)


def _add_body_paragraph(doc, text: str, size: float = design.SIZE_BODY, color_hex: Optional[str] = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(design.PARA_SPACE_AFTER_PT)
    p.paragraph_format.line_spacing = 1.2
    _write_spans(p, parse_inline(text), size, color_hex)
    return p


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------


def _build_cover(doc, model: ModelDocument, generation_date: str, baseline_commit: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(design.PAGE_WIDTH_IN)
    section.page_height = Inches(design.PAGE_HEIGHT_IN)
    section.top_margin = Inches(design.MARGIN_COVER_TOP_IN)
    section.bottom_margin = Inches(design.MARGIN_COVER_BOTTOM_IN)
    section.left_margin = Inches(design.MARGIN_COVER_SIDE_IN)
    section.right_margin = Inches(design.MARGIN_COVER_SIDE_IN)
    _unlink_header_footer(section)

    usable_w = design.PAGE_WIDTH_IN - 2 * design.MARGIN_COVER_SIDE_IN
    usable_h = design.PAGE_HEIGHT_IN - design.MARGIN_COVER_TOP_IN - design.MARGIN_COVER_BOTTOM_IN

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_no_borders(table)
    _fixed_layout(table)
    table.columns[0].width = Inches(usable_w)
    cell = table.cell(0, 0)
    cell.width = Inches(usable_w)
    _shade_cell(cell, design.OBSIDIAN_BLACK)
    tr = table.rows[0]
    trPr = tr._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(usable_h * 1440)))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

    cell.paragraphs[0].text = ""
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)

    def cover_line(text: str, size: float, color_hex: str, bold: bool = False, italic: bool = False,
                   space_before: float = 0, space_after: float = 6, font: str = design.DISPLAY_FONT):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = font
        run.font.color.rgb = _rgb(color_hex)
        return p

    # vertical spacer to push the title block toward the visual center
    cover_line("", 1, design.IVORY_WHITE, space_after=0)
    for _ in range(6):
        cover_line("", 6, design.IVORY_WHITE, space_after=0)

    cover_line("DarkSage", design.SIZE_COVER_TITLE, design.HIGHLIGHT_GOLD, bold=True, space_after=4)
    cover_line(model.title.split("—", 1)[-1].strip() or model.title, 18, design.IVORY_WHITE, space_after=18)
    cover_line(f"Version {model.version} — {model.status}", 12, design.SOFT_GRAY, space_after=4)
    cover_line(f"Owner: {model.owner}", 11, design.SOFT_GRAY, space_after=4)
    cover_line(f"Classification: {model.classification}", 11, design.SOFT_GRAY, space_after=24)
    cover_line(design.MOTTO, 14, design.SAGE_GOLD, italic=True, space_after=36, font=design.BODY_FONT)

    for _ in range(4):
        cover_line("", 6, design.IVORY_WHITE, space_after=0)

    cover_line(f"Source repository: {model.repository}", 9, design.SOFT_GRAY, space_after=2, font=design.MONO_FONT)
    cover_line(f"Source baseline commit: {baseline_commit}", 9, design.SOFT_GRAY, space_after=2, font=design.MONO_FONT)
    cover_line(f"Generation date: {generation_date}", 9, design.SOFT_GRAY, space_after=2, font=design.MONO_FONT)
    cover_line(f"Document ID: {model.doc_id}", 9, design.SOFT_GRAY, space_after=2, font=design.MONO_FONT)


# ---------------------------------------------------------------------------
# Header / footer for the body section
# ---------------------------------------------------------------------------


def _build_running_header_footer(section, short_title: str, doc_id: str, classification: str,
                                  version: str, status: str, baseline_commit: str) -> None:
    _unlink_header_footer(section)

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(2)
    tabs = hp.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(design.PAGE_WIDTH_IN - design.MARGIN_NORMAL_INSIDE_IN - design.MARGIN_NORMAL_OUTSIDE_IN))
    run = hp.add_run(f"DarkSage — {short_title}")
    run.font.size = Pt(design.SIZE_CAPTION)
    run.font.name = design.BODY_FONT
    run.font.color.rgb = _rgb(design.NEAR_OBSIDIAN_BODY_TEXT)
    hp.add_run("\t")
    r2 = hp.add_run(doc_id)
    r2.font.size = Pt(design.SIZE_CAPTION)
    r2.font.name = design.MONO_FONT
    r2.font.color.rgb = _rgb(design.SAGE_GOLD)
    r2.bold = True
    _paragraph_bottom_rule(hp, design.SAGE_GOLD, sz="6")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Page ")
    r.font.size = Pt(design.SIZE_CAPTION)
    r.font.color.rgb = _rgb(design.SOFT_GRAY)
    _add_field(fp, "PAGE", "1")
    r = fp.add_run(" of ")
    r.font.size = Pt(design.SIZE_CAPTION)
    r.font.color.rgb = _rgb(design.SOFT_GRAY)
    _add_field(fp, "NUMPAGES", "1")
    for run in fp.runs:
        run.font.size = Pt(design.SIZE_CAPTION)
        run.font.name = design.BODY_FONT
        run.font.color.rgb = _rgb(design.SOFT_GRAY)

    fp2 = footer.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp2.add_run(f"v{version} — {status}  |  Classification: {classification}")
    r.font.size = Pt(7.5)
    r.font.name = design.BODY_FONT
    r.font.color.rgb = _rgb(design.SOFT_GRAY)

    fp3 = footer.add_paragraph()
    fp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp3.add_run(f"Generated from {doc_id} v{version} — baseline {baseline_commit}")
    r.font.size = Pt(7)
    r.font.name = design.MONO_FONT
    r.font.color.rgb = _rgb(design.SOFT_GRAY)


def _new_body_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Inches(design.PAGE_WIDTH_IN)
    section.page_height = Inches(design.PAGE_HEIGHT_IN)
    section.top_margin = Inches(design.MARGIN_NORMAL_TOP_IN)
    section.bottom_margin = Inches(design.MARGIN_NORMAL_BOTTOM_IN)
    section.left_margin = Inches(design.MARGIN_NORMAL_INSIDE_IN)
    section.right_margin = Inches(design.MARGIN_NORMAL_OUTSIDE_IN)
    return section


# ---------------------------------------------------------------------------
# Document Control / Revision History / TOC dedicated pages
# ---------------------------------------------------------------------------


def _add_kv_table(doc, rows: List[List[str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_borders(table, design.STEEL_GRAY, sz="4")
    _fixed_layout(table)
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.5)
    for key, value in rows:
        row = table.add_row()
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.5)
        _shade_cell(row.cells[0], design.CHARCOAL_GRAY)
        kp = row.cells[0].paragraphs[0]
        kr = kp.add_run(key)
        kr.bold = True
        kr.font.size = Pt(design.SIZE_TABLE)
        kr.font.name = design.BODY_FONT
        kr.font.color.rgb = _rgb(design.IVORY_WHITE)
        vp = row.cells[1].paragraphs[0]
        _write_spans(vp, parse_inline(value), design.SIZE_TABLE)


def _build_document_control_page(doc, model: ModelDocument) -> None:
    h = doc.add_heading("Document Control", level=1)
    _style_heading(h, 1)
    order = [
        "Document ID", "Title", "Version", "Status", "Owner", "Contributors",
        "Classification", "Repository", "Created", "Last Updated",
        "Source Baseline Commit", "Controlling Sources", "Authority Boundary",
        "Publication Relationship",
    ]
    rows = [[k, model.doc_control[k]] for k in order if model.doc_control.get(k)]
    for k, v in model.doc_control.items():
        if k not in order:
            rows.append([k, v])
    _add_kv_table(doc, rows)
    _add_body_paragraph(
        doc,
        "Status lifecycle: Draft → Under Review → Approved → Superseded/Deprecated.",
        size=design.SIZE_CAPTION,
        color_hex=design.SOFT_GRAY,
    )


def _build_revision_history_page(doc, model: ModelDocument) -> None:
    h = doc.add_heading("Revision History", level=1)
    _style_heading(h, 1)
    table = doc.add_table(rows=1, cols=4)
    _table_borders(table, design.STEEL_GRAY, sz="4")
    _fixed_layout(table)
    widths = [0.7, 0.9, 1.6, 3.1]
    for i, w in enumerate(widths):
        table.columns[i].width = Inches(w)
    hdr = table.rows[0].cells
    for i, label in enumerate(["Version", "Date", "Author", "Summary"]):
        hdr[i].width = Inches(widths[i])
        _shade_cell(hdr[i], design.OBSIDIAN_BLACK)
        rp = hdr[i].paragraphs[0]
        rr = rp.add_run(label)
        rr.bold = True
        rr.font.color.rgb = _rgb(design.HIGHLIGHT_GOLD)
        rr.font.size = Pt(design.SIZE_TABLE_HEADER)
        rr.font.name = design.BODY_FONT
    for rev in model.revision_history:
        row = table.add_row()
        cells = row.cells
        for i, w in enumerate(widths):
            cells[i].width = Inches(w)
        for i, val in enumerate([rev.version, rev.date, rev.author, rev.summary]):
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(4)
            _write_spans(cells[i].paragraphs[0], parse_inline(val), design.SIZE_TABLE)


def _build_toc_page(doc) -> None:
    h = doc.add_heading("Table of Contents", level=1)
    _style_heading(h, 1)
    note = doc.add_paragraph()
    note_run = note.add_run(
        "This Table of Contents is a Word/LibreOffice field. If it appears empty, "
        "select it and press F9 (or right-click → Update Field) to populate it "
        "from this document's own headings."
    )
    note_run.italic = True
    note_run.font.size = Pt(design.SIZE_CAPTION)
    note_run.font.color.rgb = _rgb(design.SOFT_GRAY)
    p = doc.add_paragraph()
    _add_field(p, 'TOC \\o "1-2" \\h \\z \\u', "Right-click and choose Update Field to generate the Table of Contents.")


# ---------------------------------------------------------------------------
# Heading style setup
# ---------------------------------------------------------------------------


def _style_heading(paragraph, model_level: int, source_level: int = 0) -> None:
    """model_level: 1 -> Word Heading 1 (top TOC level), 2 -> Word Heading 2,
    3 -> Word Heading 3 (used only for markdown H4; H1-H3 behavior below is
    unchanged from before H4 support was added).

    source_level: the original markdown heading level (2, 3, or 4), used only
    to select the H4-specific style below without altering the existing
    H1/else behavior that H2 and H3 both still receive identically."""
    for run in paragraph.runs:
        run.font.name = design.HEADING_FONT
        run.font.color.rgb = _rgb(design.NEAR_OBSIDIAN_BODY_TEXT)
    if model_level == 1:
        for run in paragraph.runs:
            run.font.size = Pt(design.SIZE_H1)
            run.font.color.rgb = _rgb(design.OBSIDIAN_BLACK)
        paragraph.paragraph_format.space_before = Pt(design.H1_SPACE_BEFORE_PT)
        paragraph.paragraph_format.space_after = Pt(design.H1_SPACE_AFTER_PT)
        _paragraph_bottom_rule(paragraph, design.SAGE_GOLD, sz="16")
    elif source_level == 4:
        for run in paragraph.runs:
            run.font.size = Pt(design.SIZE_H4)
            run.font.italic = True
        paragraph.paragraph_format.space_before = Pt(design.H2_SPACE_BEFORE_PT * 0.6)
        paragraph.paragraph_format.space_after = Pt(design.H2_SPACE_AFTER_PT * 0.6)
    else:
        for run in paragraph.runs:
            run.font.size = Pt(design.SIZE_H2)
        paragraph.paragraph_format.space_before = Pt(design.H2_SPACE_BEFORE_PT)
        paragraph.paragraph_format.space_after = Pt(design.H2_SPACE_AFTER_PT)
        _paragraph_bottom_rule(paragraph, design.SAGE_GOLD, sz="8")


# ---------------------------------------------------------------------------
# Body rendering
# ---------------------------------------------------------------------------


def _render_table(doc, block: Table) -> None:
    ncols = len(block.header)
    table = doc.add_table(rows=1, cols=ncols)
    _table_borders(table, design.STEEL_GRAY, sz="4")
    _fixed_layout(table)
    body_width = design.PAGE_WIDTH_IN - design.MARGIN_NORMAL_INSIDE_IN - design.MARGIN_NORMAL_OUTSIDE_IN
    col_w = body_width / ncols
    for i in range(ncols):
        table.columns[i].width = Inches(col_w)
    hdr_cells = table.rows[0].cells
    for i, label in enumerate(block.header):
        hdr_cells[i].width = Inches(col_w)
        _shade_cell(hdr_cells[i], design.CHARCOAL_GRAY)
        rp = hdr_cells[i].paragraphs[0]
        rr = rp.add_run(label)
        rr.bold = True
        rr.font.color.rgb = _rgb(design.IVORY_WHITE)
        rr.font.size = Pt(design.SIZE_TABLE_HEADER)
        rr.font.name = design.BODY_FONT
    # Wide tables (5+ cols or long cell text) step the font size down per
    # DSF-003 SS4.8 rather than truncating any column.
    font_size = design.SIZE_TABLE if ncols <= 4 else max(7.5, design.SIZE_TABLE - (ncols - 4) * 0.5)
    for row_vals in block.rows:
        row = table.add_row()
        cells = row.cells
        for i in range(ncols):
            val = row_vals[i] if i < len(row_vals) else ""
            cells[i].width = Inches(col_w)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(3)
            _write_spans(cells[i].paragraphs[0], parse_inline(val), font_size)


def _render_figure(doc, block: Figure, registry: DiagramRegistry) -> None:
    h = doc.add_heading(f"Figure {block.number} — {block.title}", level=2)
    _style_heading(h, 2)
    png = registry.png_path(block.number)
    if png and png.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        body_width = design.PAGE_WIDTH_IN - design.MARGIN_NORMAL_INSIDE_IN - design.MARGIN_NORMAL_OUTSIDE_IN
        shape = run.add_picture(str(png), width=Inches(min(6.0, body_width)))
        alt = registry.accessibility_text(block.number) or block.caption
        _set_alt_text(shape, alt)
    else:
        note = doc.add_paragraph()
        nr = note.add_run(
            f"[Figure {block.number} not yet rendered — see docs/publication/DIAGRAM_REGISTER.md row {block.number}.]"
        )
        nr.italic = True
        nr.font.color.rgb = _rgb(design.WARNING_AMBER)
        nr.font.size = Pt(design.SIZE_CAPTION)
    if block.caption:
        _add_body_paragraph(doc, block.caption, size=design.SIZE_CAPTION, color_hex=design.SOFT_GRAY)
    alt_text = registry.accessibility_text(block.number)
    if alt_text:
        ap = doc.add_paragraph()
        label = ap.add_run("Accessibility description: ")
        label.bold = True
        label.font.size = Pt(design.SIZE_CAPTION)
        label.font.color.rgb = _rgb(design.STEEL_GRAY)
        desc = ap.add_run(alt_text)
        desc.italic = True
        desc.font.size = Pt(design.SIZE_CAPTION)
        desc.font.color.rgb = _rgb(design.SOFT_GRAY)


def _render_blockquote(doc, block: Blockquote) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(design.PARA_SPACE_AFTER_PT)
    _paragraph_bottom_rule(p, design.SAGE_GOLD, sz="0")
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), design.SAGE_GOLD)
    pBdr.append(left)
    _write_spans(p, parse_inline(block.text), design.SIZE_BODY, color_hex=design.STEEL_GRAY)
    for run in p.runs:
        run.italic = True


def _render_code_block(doc, block: CodeBlock) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(design.PARA_SPACE_AFTER_PT)
    for i, line in enumerate(block.text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = design.MONO_FONT
        r.font.size = Pt(design.SIZE_MONO - 1)
        r.font.color.rgb = _rgb(design.NEAR_OBSIDIAN_BODY_TEXT)


def _render_list(doc, block: ListBlock) -> None:
    style = "List Number" if block.ordered else "List Bullet"
    for item in block.items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        _write_spans(p, parse_inline(item), design.SIZE_BODY)


def _render_body_blocks(doc, model: ModelDocument, registry: DiagramRegistry) -> None:
    # An H3 that appears before any real (non-skipped) H2 -- e.g. "Document
    # ID Rationale" nested under the skipped "Document Control" H2, whose
    # table/heading are pulled onto the dedicated Document-Control page
    # instead of repeating in the body -- would otherwise render as a
    # sub-heading with no visible parent in the flowing body/TOC. Promoting
    # it to a top-level heading keeps its real content (never dropped)
    # without an orphaned-looking indent level.
    seen_real_h2 = False
    for block in model.blocks:
        if isinstance(block, Heading):
            if block.skip_in_body:
                continue
            if block.level == 2:
                seen_real_h2 = True
                level = 1
            elif block.level == 4:
                level = 3
            else:
                level = 1 if not seen_real_h2 else 2
            h = doc.add_heading("", level=level)
            size = design.SIZE_H4 if block.level == 4 else (design.SIZE_H1 if level == 1 else design.SIZE_H2)
            _write_spans(h, parse_inline(block.text), size)
            _style_heading(h, level, source_level=block.level)
        elif isinstance(block, Paragraph):
            if block.skip_in_body:
                continue
            _add_body_paragraph(doc, block.text)
        elif isinstance(block, Table):
            if block.skip_in_body:
                continue
            _render_table(doc, block)
        elif isinstance(block, Figure):
            _render_figure(doc, block, registry)
        elif isinstance(block, Blockquote):
            _render_blockquote(doc, block)
        elif isinstance(block, CodeBlock):
            _render_code_block(doc, block)
        elif isinstance(block, ListBlock):
            _render_list(doc, block)
        elif isinstance(block, Rule):
            p = doc.add_paragraph()
            _paragraph_bottom_rule(p, design.SAGE_GOLD, sz="6")


# ---------------------------------------------------------------------------
# Top-level entry point
# ---------------------------------------------------------------------------


def render_docx(
    model: ModelDocument,
    out_path: Path,
    repo_root: Path,
    generation_date: str,
    baseline_commit: str,
    short_title: Optional[str] = None,
) -> None:
    registry = DiagramRegistry(repo_root)
    doc = DocxDocument()

    normal = doc.styles["Normal"]
    normal.font.name = design.BODY_FONT
    normal.font.size = Pt(design.SIZE_BODY)

    short_title = short_title or model.title

    _build_cover(doc, model, generation_date, baseline_commit)

    body_section = _new_body_section(doc)
    _build_running_header_footer(
        body_section, short_title, model.doc_id, model.classification,
        model.version, model.status, baseline_commit,
    )

    _build_document_control_page(doc, model)
    doc.add_page_break()
    _build_revision_history_page(doc, model)
    doc.add_page_break()
    _build_toc_page(doc)
    doc.add_page_break()

    _render_body_blocks(doc, model, registry)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    _normalize_zip_timestamps(out_path)


_FIXED_ZIP_DATE = (1980, 1, 1, 0, 0, 0)  # DOS-epoch floor; classic reproducible-build convention


def _normalize_zip_timestamps(path: Path) -> None:
    """python-docx's OPC writer (docx/opc/phys_pkg.py) calls
    zipfile.ZipFile.writestr(name, blob) with a bare name, which stamps
    every zip entry with the wall-clock time at save() -- the one source of
    non-determinism in an otherwise byte-for-byte reproducible DOCX. This
    rewrites the archive with every entry's date_time pinned to a fixed
    constant, in the same order and with the same compression, so two
    generation runs against identical input produce byte-identical output."""
    with zipfile.ZipFile(path, "r") as zin:
        items = [(info, zin.read(info.filename)) for info in zin.infolist()]
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for info, data in items:
            fixed = zipfile.ZipInfo(info.filename, date_time=_FIXED_ZIP_DATE)
            fixed.compress_type = info.compress_type
            fixed.external_attr = info.external_attr
            zout.writestr(fixed, data)
