"""Unit tests for scripts/publication/docgen/ (the DOCX/PDF generation
pipeline). Mirrors test_publication_tools.py's fixture style: every test
uses a throwaway tempfile.TemporaryDirectory tree, never the real repo, and
never depends on wall-clock time or this machine's real Git state beyond
what it explicitly supplies.
"""

from __future__ import annotations

import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docgen.docx_render import render_docx  # noqa: E402
from docgen.model import (  # noqa: E402
    Blockquote,
    CodeBlock,
    Figure,
    Heading,
    ListBlock,
    Paragraph,
    Table,
)
from docgen.parser import parse_markdown  # noqa: E402
from docgen.pdf_render import render_pdf  # noqa: E402
from docx import Document as DocxDocument  # noqa: E402

_SAMPLE_MD = """# DSF-999 — Sample Test Document

## 1. Document Control

| Field | Value |
|---|---|
| Document ID | DSF-999 |
| Title | Sample Test Document |
| Version | 0.1.0 |
| Status | Draft |
| Owner | Test Owner |
| Classification | Internal |
| Repository | LightSwornPVP/DarkSage |

## Revision History

| Version | Date | Author | Summary |
|---|---|---|---|
| 0.1.0 | 2026-07-25 | Test Owner | First draft. |

## 2. Introduction

This is a **bold** claim with `DS-001` inline code and *italic* emphasis.

| ID | Title | Class. |
|---|---|---|
| DS-TEST-001 | Example Requirement | **Committed/MVP** |
| DS-TEST-002 | Another Requirement | Planned |

- First bullet
- Second bullet

1. Step one
2. Step two

> A quoted governance statement.

```
plain code block
second line
```

**Figure 42 — A Sample Figure** *(placeholder; not yet rendered)*. Describes something.

### 2.1 Subsection

More prose here for the subsection.

#### DS-BL-999 — Sample H4 Backlog Entry

A fourth-level heading, representative of DS-006/DS-013/DS-014's real heading pattern at that depth.
"""

_H4_REPRESENTATIVE_MD = """# DS-999 — Representative H4 Sample

## Document Control

| Field | Value |
|---|---|
| Document ID | DS-999 |
| Title | Representative H4 Sample |
| Version | 0.1.0 |
| Status | Draft |
| Owner | Test Owner |
| Classification | Internal |
| Repository | LightSwornPVP/DarkSage |

## Revision History

| Version | Date | Author | Summary |
|---|---|---|---|
| 0.1.0 | 2026-07-25 | Test Owner | First draft. |

## 1. Section

### 1.1 Subsection

#### DS-IDEA-002 — Market Tokenization Inspired by Kronos

Exploratory research idea, representative of DS-014's actual H4 pattern.

#### DS-BL-017 — Sage Memory (Persistent, Cross-Session)

Backlog entry, representative of DS-013's actual H4 pattern.
"""


class TestMarkdownParsing(unittest.TestCase):
    def setUp(self):
        self.doc = parse_markdown(_SAMPLE_MD, source_path="sample.md")

    def test_title_and_doc_id(self):
        self.assertIn("Sample Test Document", self.doc.title)
        self.assertEqual(self.doc.doc_id, "DSF-999")

    def test_doc_control_fields_captured(self):
        self.assertEqual(self.doc.doc_control["Version"], "0.1.0")
        self.assertEqual(self.doc.doc_control["Status"], "Draft")
        self.assertEqual(self.doc.doc_control["Owner"], "Test Owner")

    def test_revision_history_captured(self):
        self.assertEqual(len(self.doc.revision_history), 1)
        rev = self.doc.revision_history[0]
        self.assertEqual(rev.version, "0.1.0")
        self.assertEqual(rev.author, "Test Owner")
        self.assertEqual(rev.summary, "First draft.")

    def test_doc_control_and_revision_tables_flagged_skip_in_body(self):
        tables = [b for b in self.doc.blocks if isinstance(b, Table)]
        # doc-control table (3 cols? no -- 2 cols) and revision-history table (4 cols)
        skip_flagged = [t for t in tables if t.skip_in_body]
        self.assertEqual(len(skip_flagged), 2)

    def test_body_table_not_skipped_and_has_expected_rows(self):
        tables = [b for b in self.doc.blocks if isinstance(b, Table) and not b.skip_in_body]
        self.assertEqual(len(tables), 1)
        body_table = tables[0]
        self.assertEqual(body_table.header, ["ID", "Title", "Class."])
        self.assertEqual(len(body_table.rows), 2)
        self.assertEqual(body_table.rows[0][0], "DS-TEST-001")

    def test_figure_marker_extracted(self):
        figures = [b for b in self.doc.blocks if isinstance(b, Figure)]
        self.assertEqual(len(figures), 1)
        self.assertEqual(figures[0].number, 42)
        self.assertEqual(figures[0].title, "A Sample Figure")

    def test_list_blocks_extracted(self):
        lists = [b for b in self.doc.blocks if isinstance(b, ListBlock)]
        self.assertEqual(len(lists), 2)
        bullet, ordered = lists[0], lists[1]
        self.assertFalse(bullet.ordered)
        self.assertEqual(bullet.items, ["First bullet", "Second bullet"])
        self.assertTrue(ordered.ordered)
        self.assertEqual(ordered.items, ["Step one", "Step two"])

    def test_blockquote_extracted(self):
        quotes = [b for b in self.doc.blocks if isinstance(b, Blockquote)]
        self.assertEqual(len(quotes), 1)
        self.assertEqual(quotes[0].text, "A quoted governance statement.")

    def test_code_block_extracted(self):
        codes = [b for b in self.doc.blocks if isinstance(b, CodeBlock)]
        self.assertEqual(len(codes), 1)
        self.assertIn("plain code block", codes[0].text)
        self.assertIn("second line", codes[0].text)

    def test_headings_levels_and_anchors_unique(self):
        headings = [b for b in self.doc.blocks if isinstance(b, Heading)]
        levels = [h.level for h in headings]
        self.assertIn(2, levels)
        self.assertIn(3, levels)
        anchors = [h.anchor for h in headings]
        self.assertEqual(len(anchors), len(set(anchors)))

    def test_plain_paragraph_captured(self):
        paras = [b for b in self.doc.blocks if isinstance(b, Paragraph)]
        joined = " ".join(p.text for p in paras)
        self.assertIn("bold", joined)
        self.assertIn("DS-001", joined)


class RenderFixtureTestCase(unittest.TestCase):
    """Renders the sample document into a scratch directory that also
    stands in as an (empty) repo root -- DiagramRegistry gracefully finds
    no rendered diagrams there, exercising the "not yet rendered" fallback
    path used by real not-yet-authored figures."""

    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.root = Path(self._tmp.name)
        self.model = parse_markdown(_SAMPLE_MD, source_path="sample.md")

    def tearDown(self):
        self._tmp.cleanup()


class TestDocxStructuralValidity(RenderFixtureTestCase):
    def test_docx_is_a_valid_zip_and_reopenable(self):
        out = self.root / "out.docx"
        render_docx(self.model, out, self.root, "2026-07-25T00:00:00+00:00", "deadbeef", short_title="Sample")
        self.assertTrue(zipfile.is_zipfile(out))
        reopened = DocxDocument(str(out))
        self.assertGreater(len(reopened.paragraphs), 0)
        self.assertGreaterEqual(len(reopened.sections), 2)  # cover section + body section
        self.assertGreaterEqual(len(reopened.tables), 3)  # doc control + revision history + body table

    def test_docx_contains_figure_placeholder_note_when_not_rendered(self):
        out = self.root / "out.docx"
        render_docx(self.model, out, self.root, "2026-07-25T00:00:00+00:00", "deadbeef", short_title="Sample")
        reopened = DocxDocument(str(out))
        all_text = "\n".join(p.text for p in reopened.paragraphs)
        self.assertIn("Figure 42", all_text)
        self.assertIn("not yet rendered", all_text)


class TestPdfStructuralValidity(RenderFixtureTestCase):
    def test_pdf_starts_with_pdf_header(self):
        out = self.root / "out.pdf"
        pages = render_pdf(self.model, out, self.root, "2026-07-25T00:00:00+00:00", "deadbeef", short_title="Sample")
        with open(out, "rb") as fh:
            head = fh.read(5)
        self.assertEqual(head, b"%PDF-")
        self.assertGreater(pages, 0)

    def test_pdf_page_count_matches_reported_count(self):
        out = self.root / "out.pdf"
        pages = render_pdf(self.model, out, self.root, "2026-07-25T00:00:00+00:00", "deadbeef", short_title="Sample")
        raw = out.read_bytes()
        # Every page object in a reportlab-produced PDF declares /Type /Page
        # (not /Pages, the tree root node) -- a reliable, dependency-free
        # page count. Negative lookahead excludes the "/Type /Pages" root.
        import re

        counted = len(re.findall(rb"/Type\s*/Page(?!s)", raw))
        self.assertEqual(counted, pages)


class TestDeterministicGeneration(RenderFixtureTestCase):
    def test_docx_identical_across_two_runs(self):
        out1 = self.root / "run1.docx"
        out2 = self.root / "run2.docx"
        render_docx(self.model, out1, self.root, "2026-07-25T00:00:00+00:00", "deadbeef", short_title="Sample")
        render_docx(self.model, out2, self.root, "2026-07-25T00:00:00+00:00", "deadbeef", short_title="Sample")
        self.assertEqual(out1.read_bytes(), out2.read_bytes())

    def test_pdf_identical_across_two_runs(self):
        out1 = self.root / "run1.pdf"
        out2 = self.root / "run2.pdf"
        render_pdf(self.model, out1, self.root, "2026-07-25T00:00:00+00:00", "deadbeef", short_title="Sample")
        render_pdf(self.model, out2, self.root, "2026-07-25T00:00:00+00:00", "deadbeef", short_title="Sample")
        self.assertEqual(out1.read_bytes(), out2.read_bytes())

    def test_different_baseline_commit_changes_output(self):
        out1 = self.root / "run1.pdf"
        out2 = self.root / "run2.pdf"
        render_pdf(self.model, out1, self.root, "2026-07-25T00:00:00+00:00", "deadbeef", short_title="Sample")
        render_pdf(self.model, out2, self.root, "2026-07-25T00:00:00+00:00", "cafef00d", short_title="Sample")
        self.assertNotEqual(out1.read_bytes(), out2.read_bytes())


class TestH4HeadingSupport(unittest.TestCase):
    """Regression coverage for the H1 finding: H4 (####) headings were
    previously unsupported by the parser and rendered as literal text in
    generated DOCX/PDF output."""

    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.root = Path(self._tmp.name)
        self.model = parse_markdown(_H4_REPRESENTATIVE_MD, source_path="sample.md")

    def tearDown(self):
        self._tmp.cleanup()

    def test_parser_recognizes_h4_as_level_4(self):
        headings = [b for b in self.model.blocks if isinstance(b, Heading)]
        levels = [h.level for h in headings]
        self.assertIn(4, levels)
        h4 = [h for h in headings if h.level == 4]
        self.assertEqual(len(h4), 2)
        self.assertEqual(h4[0].text, "DS-IDEA-002 — Market Tokenization Inspired by Kronos")
        self.assertEqual(h4[1].text, "DS-BL-017 — Sage Memory (Persistent, Cross-Session)")

    def test_h1_h2_h3_levels_unaffected_by_h4_support(self):
        headings = [b for b in self.model.blocks if isinstance(b, Heading)]
        levels = sorted({h.level for h in headings})
        self.assertEqual(levels, [2, 3, 4])
        self.assertEqual(self.model.title, "DS-999 — Representative H4 Sample")

    def test_docx_h4_text_present_and_not_literal(self):
        out = self.root / "h4.docx"
        render_docx(self.model, out, self.root, "2026-07-25T00:00:00+00:00", "deadbeef", short_title="Sample")
        reopened = DocxDocument(str(out))
        all_text = "\n".join(p.text for p in reopened.paragraphs)
        self.assertIn("DS-IDEA-002", all_text)
        self.assertIn("DS-BL-017", all_text)
        self.assertNotIn("####", all_text)

    def test_pdf_h4_text_present_and_not_literal(self):
        out = self.root / "h4.pdf"
        render_pdf(self.model, out, self.root, "2026-07-25T00:00:00+00:00", "deadbeef", short_title="Sample")
        # PDF text isn't trivially extractable without an extra dependency;
        # assert against the raw content streams instead (reportlab emits
        # heading text as literal Tj/TJ string operands).
        raw = out.read_bytes()
        self.assertIn(b"DS-IDEA-002", raw)
        self.assertNotIn(b"####", raw)

    def test_sample_fixture_h4_also_has_no_literal_hashes(self):
        # Exercises the shared _SAMPLE_MD fixture's own H4 line (used by the
        # rest of this file's tests) for the same literal-#### regression.
        model = parse_markdown(_SAMPLE_MD, source_path="sample.md")
        out = self.root / "sample_h4.docx"
        render_docx(model, out, self.root, "2026-07-25T00:00:00+00:00", "deadbeef", short_title="Sample")
        reopened = DocxDocument(str(out))
        all_text = "\n".join(p.text for p in reopened.paragraphs)
        self.assertIn("DS-BL-999", all_text)
        self.assertNotIn("####", all_text)


if __name__ == "__main__":
    unittest.main()
